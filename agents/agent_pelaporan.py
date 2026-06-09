import logging
import os
from datetime import date
from docx import Document

logger = logging.getLogger(__name__)

RISK_EMOJI = {"High": "🔴", "Medium": "🟡", "Low": "🟢"}


class AgentPelaporan:
    def __init__(self, reports_dir: str = "reports"):
        self.reports_dir = reports_dir

    def run(self, dataset: dict) -> dict:
        """Generate laporan markdown dan simpan file .docx."""
        report_date = dataset.get("date", str(date.today()))
        findings = dataset.get("findings", [])
        summary = dataset.get("summary", {})

        logger.info("Agent Pelaporan: membuat laporan untuk tanggal %s dengan %d findings", report_date, len(findings))

        markdown = self._build_markdown(report_date, findings, summary)
        docx_path = self._save_docx(report_date, findings, summary)

        logger.info("Agent Pelaporan: laporan selesai, docx disimpan di %s", docx_path)
        return {"markdown": markdown, "docx_path": docx_path, "date": report_date}

    def _build_markdown(self, report_date: str, findings: list, summary: dict) -> str:
        lines = [
            "# 📋 Laporan Harian Gurimbang Safety",
            f"**Tanggal:** {report_date}",
            "",
            "---",
            "",
            "## Ringkasan Eksekutif",
            "",
            "| Indikator | Jumlah |",
            "|-----------|--------|",
            f"| Total Temuan Open | {summary.get('total', 0)} |",
            f"| Overdue | {summary.get('overdue', 0)} |",
            f"| High Risk | {summary.get('high_risk', 0)} |",
            f"| Blindspot | {summary.get('blindspot', 0)} |",
            f"| Perlu Intervensi | {summary.get('needs_intervention', 0)} |",
            "",
            "---",
            "",
            "## Daftar Intervensi",
            "",
            "| # | Lokasi | Temuan | Risiko | Due Date | Status | PJ Area | Kontak |",
            "|---|--------|--------|--------|----------|--------|---------|--------|",
        ]

        for i, f in enumerate(findings, 1):
            risk = f.get("risk_level", "-")
            emoji = RISK_EMOJI.get(risk, "⚪")
            overdue_flag = " ⚠️ OVERDUE" if f.get("is_overdue") else ""
            blindspot_flag = " 🔍 BLINDSPOT" if f.get("is_blindspot") else ""
            lines.append(
                f"| {i} "
                f"| {f.get('location_name', '-')} ({f.get('site', '-')}) "
                f"| {str(f.get('description', '-'))[:60]} "
                f"| {emoji} {risk} "
                f"| {f.get('due_date', '-')}{overdue_flag} "
                f"| {f.get('status', '-')}{blindspot_flag} "
                f"| {f.get('pj_name', '-')} "
                f"| {f.get('pj_contact', '-')} |"
            )

        lines += [
            "",
            "---",
            "",
            "*Laporan dibuat otomatis oleh Gurimbang Safety Intelligence System*",
        ]
        return "\n".join(lines)

    def _save_docx(self, report_date: str, findings: list, summary: dict) -> str:
        out_dir = os.path.join(self.reports_dir, report_date)
        os.makedirs(out_dir, exist_ok=True)
        out_path = os.path.join(out_dir, f"laporan-safety-{report_date}.docx")

        doc = Document()
        doc.add_heading("Laporan Harian Gurimbang Safety", 0)
        doc.add_paragraph(f"Tanggal: {report_date}")

        doc.add_heading("Ringkasan Eksekutif", level=1)
        tbl_summary = doc.add_table(rows=1, cols=2)
        tbl_summary.style = "Table Grid"
        tbl_summary.rows[0].cells[0].text = "Indikator"
        tbl_summary.rows[0].cells[1].text = "Jumlah"
        for label, key in [
            ("Total Temuan Open", "total"),
            ("Overdue", "overdue"),
            ("High Risk", "high_risk"),
            ("Blindspot", "blindspot"),
            ("Perlu Intervensi", "needs_intervention"),
        ]:
            row = tbl_summary.add_row().cells
            row[0].text = label
            row[1].text = str(summary.get(key, 0))

        doc.add_heading("Daftar Intervensi", level=1)
        if findings:
            cols = ["No", "Lokasi", "Temuan", "Risiko", "Due Date", "Status", "PJ Area", "Kontak"]
            tbl = doc.add_table(rows=1, cols=len(cols))
            tbl.style = "Table Grid"
            for i, col in enumerate(cols):
                tbl.rows[0].cells[i].text = col
            for idx, f in enumerate(findings, 1):
                cells = tbl.add_row().cells
                cells[0].text = str(idx)
                cells[1].text = f"{f.get('location_name', '-')} ({f.get('site', '-')})"
                cells[2].text = str(f.get("description", "-"))[:80]
                cells[3].text = f.get("risk_level", "-")
                due = f.get("due_date", "-")
                cells[4].text = f"{due} ⚠️" if f.get("is_overdue") else str(due)
                cells[5].text = f.get("status", "-")
                cells[6].text = f.get("pj_name", "-")
                cells[7].text = f.get("pj_contact", "-")

        doc.save(out_path)
        return out_path
